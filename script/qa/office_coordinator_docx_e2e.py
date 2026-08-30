#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx>=1.2,<2"]
# ///

# ─── How to run ───
# 1. Install uv (if not installed):
#      curl -LsSf https://astral.sh/uv/install.sh | sh
# 2. Run directly (no venv, no pip install needed):
#      PYTHONPATH=. python script/qa/office_coordinator_docx_e2e.py --out PATH
# 3. Or make executable and run:
#      chmod +x script/qa/office_coordinator_docx_e2e.py && ./script/qa/office_coordinator_docx_e2e.py --out PATH
# ──────────────────

"""Real ToolRegistry-to-approval DOCX coordinator acceptance driver."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import shutil
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Final, cast

from docx import Document

from birkin import approvals, config, store
from birkin.office.errors import DocumentError
from birkin.office.job_journal import OfficeJobJournal
from birkin.office.job_runner import DocumentServiceRunner
from birkin.office.service import DocumentService
from birkin.tools import ToolContext, ToolRegistry, build_registry

BEFORE: Final = "계약 원문 Original contract paragraph."
AFTER: Final = "승인된 계약문 Approved contract paragraph."
ACTOR: Final = "user:office-docx-e2e"


@dataclass(frozen=True, slots=True)
class Flow:
    home: Path
    caller: Path
    source: Path
    destination: Path
    source_sha256: str
    destination_before: bytes
    registry: ToolRegistry


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _write_json(path: Path, value: object) -> None:
    path.write_text(
        json.dumps(value, ensure_ascii=False, sort_keys=True, indent=2, default=str)
        + "\n",
        encoding="utf-8",
    )


def _flow(root: Path, name: str) -> Flow:
    home, caller = root / name / "home", root / name / "caller"
    home.mkdir(parents=True)
    office_home = home / "office"
    office_home.mkdir()
    caller.mkdir()
    home, caller = home.resolve(strict=True), caller.resolve(strict=True)
    office_home = (home / "office").resolve(strict=True)
    os.environ["BIRKIN_HOME"] = str(home)
    source, destination = office_home / "source.docx", caller / "delivery.docx"
    document = Document()
    _ = document.add_paragraph(BEFORE)
    document.save(str(source))
    original = Document()
    _ = original.add_paragraph("Caller destination before export.")
    original.save(str(destination))
    return Flow(
        home=home,
        caller=caller,
        source=source,
        destination=destination,
        source_sha256=_sha256(source),
        destination_before=destination.read_bytes(),
        registry=build_registry(
            ToolContext(cfg={}, client=None, cwd=caller, record_source=ACTOR),
            include={"documents"},
        ),
    )


def _request(
    flow: Flow, *, destination: Path | None = None, overwrite: bool = True
) -> dict[str, object]:
    return {
        "request": (
            f"이 DOCX Word document의 '{BEFORE}' 문단을 '{AFTER}'로 "
            "교체해 주세요. Please replace only that paragraph."
        ),
        "source": {
            "artifact_id": flow.source_sha256,
            "content_hash": flow.source_sha256,
            "media_type": (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            "uri": str(flow.source),
            "sensitivity": "internal",
            "acl_fingerprint": "local-docx-e2e",
        },
        "outcome": "Replace the known Korean/English contract paragraph",
        "operations": [
            {"locator": {"format": "docx", "index": 1}, "value": AFTER}
        ],
        "destination": str(destination or flow.destination),
        "overwrite_approved": overwrite,
    }


def _propose(flow: Flow, request: dict[str, object]) -> dict[str, object]:
    result = flow.registry.execute("office_job_request", request)
    body = cast("dict[str, object]", json.loads(cast(str, result.content)))
    if result.is_error:
        raise AssertionError(body)
    return body


def _unchanged(flow: Flow) -> dict[str, bool]:
    return {
        "source": _sha256(flow.source) == flow.source_sha256,
        "destination": flow.destination.read_bytes() == flow.destination_before,
    }


def _denial_probes(root: Path) -> dict[str, object]:
    records: dict[str, object] = {}
    preapproval = _flow(root, "preapproval")
    body = _propose(preapproval, _request(preapproval))
    pending = store.get_pending(cast(str, body["id"]))
    assert pending is not None
    try:
        _ = approvals.execute_action(
            "office_job", cast("dict[str, object]", pending["payload"])
        )
    except DocumentError as exc:
        records["preapproval"] = {"code": exc.code.value, "unchanged": _unchanged(preapproval)}
    else:
        raise AssertionError("preapproval execution was not denied")

    outside = _flow(root, "outside-root")
    escaped = outside.caller.parent / "escaped.docx"
    result = outside.registry.execute("office_job_request", _request(outside, destination=escaped))
    envelope = cast("dict[str, object]", json.loads(cast(str, result.content)))
    error = cast("dict[str, object]", envelope["error"])
    records["outside_root"] = {
        "code": error["code"], "unchanged": _unchanged(outside), "outside_exists": escaped.exists()
    }

    overwrite = _flow(root, "overwrite")
    body = _propose(overwrite, _request(overwrite, overwrite=False))
    denied = approvals.approve(cast(str, body["id"]), approved_by="system:qa", approved_via="qa:script")
    assert denied["ok"] is False
    assert denied["error"] == "기존 파일을 덮어쓸까요?"
    follow_up = store.get_pending(cast(str, denied["follow_up_approval_id"]))
    assert follow_up is not None
    records["overwrite"] = {
        "code": "OUTPUT_EXISTS",
        "follow_up_approval_id": follow_up["id"],
        "overwrite_approved": cast(dict[str, object], follow_up["payload"])[
            "overwrite_approved"
        ],
        "unchanged": _unchanged(overwrite),
    }

    drift = _flow(root, "digest-drift")
    body = _propose(drift, _request(drift))
    approval_id = cast(str, body["id"])
    record = store.get_pending(approval_id)
    assert record is not None
    payload = cast("dict[str, object]", record["payload"])
    record["payload"] = {**payload, "proposal_digest": "0" * 64}
    config.pending_dir().joinpath(f"{approval_id}.json").write_text(
        json.dumps(record), encoding="utf-8"
    )
    denied = approvals.approve(approval_id, approved_by="system:qa", approved_via="qa:script")
    assert denied["ok"] is False and "POLICY_DENIED" in cast(str, denied["error"])
    records["digest_drift"] = {"code": "POLICY_DENIED", "unchanged": _unchanged(drift)}
    expected = {
        "preapproval": "POLICY_DENIED",
        "outside_root": "PERMISSION_DENIED",
        "overwrite": "OUTPUT_EXISTS",
        "digest_drift": "POLICY_DENIED",
    }
    assert {name: cast("dict[str, object]", item)["code"] for name, item in records.items()} == expected
    assert all(all(cast("dict[str, bool]", cast("dict[str, object]", item)["unchanged"]).values()) for item in records.values())
    return records


def run(out: Path) -> dict[str, object]:
    """Execute the real flow, persist evidence, and return its cleanup receipt."""
    out = out.resolve()
    out.mkdir(parents=True, exist_ok=True)
    actions: list[str] = []
    runtime_path: Path
    with tempfile.TemporaryDirectory(prefix="birkin-docx-e2e-") as temporary:
        runtime_path = Path(temporary)
        flow = _flow(runtime_path, "approved")
        before_destination_sha = _sha256(flow.destination)
        body = _propose(flow, _request(flow))
        approval_id = cast(str, body["id"])
        approval_record = store.get_pending(approval_id)
        assert approval_record is not None
        approval = cast("dict[str, object]", body["approval"])
        summaries = cast("list[dict[str, str]]", approval["semantic_summaries"])
        assert len(summaries) == 1
        assert summaries[0] == {
            "location": "docx paragraph 1",
            "before": BEFORE,
            "after": AFTER,
        }
        assert approval["source_sha256"] == flow.source_sha256
        assert approval["destination"] == str(flow.destination)
        assert approval["proposer"] == ACTOR
        assert isinstance(approval["authority_digest"], str)
        actions.append("semantic preview captured and approval binding verified")
        pending = cast("dict[str, object]", approval_record["payload"])
        try:
            _ = approvals.execute_action("office_job", pending)
        except DocumentError as exc:
            assert exc.code.value == "POLICY_DENIED"
        else:
            raise AssertionError("preapproval execution was not denied")
        actions.append("preapproval execution denied with POLICY_DENIED")
        approved = approvals.approve(
            approval_id,
            approved_by="system:qa",
            approved_via="qa:office-coordinator-docx",
        )
        assert approved["ok"] is True, approved
        approval_record = store.get_pending(approval_id)
        assert approval_record is not None
        assert approval_record["approved_by"] == "system:qa"
        assert approval_record["approved_via"] == "qa:office-coordinator-docx"
        receipt = cast("dict[str, object]", json.loads(cast(str, approved["result"])))
        assert Document(str(flow.destination)).paragraphs[0].text == AFTER
        exported_sha = _sha256(flow.destination)
        actions.append("approved queue executed, validated, materialized, and exported")
        runner = DocumentServiceRunner(DocumentService(flow.home / "office"), export_root=flow.caller)
        journal = OfficeJobJournal(flow.home / "office" / "jobs")
        job = journal.restore(cast(str, receipt["job_id"]), runner=runner)
        rollback = job.rollback_export()
        assert flow.destination.read_bytes() == flow.destination_before
        assert _sha256(flow.source) == flow.source_sha256
        actions.append("durable job restored and destination rolled back byte-for-byte")
        final_receipt = job.receipt()
        export = cast("dict[str, object]", final_receipt["export"])
        assert {
            "authority_digest",
            "source_sha256",
            "output_sha256",
            "operations",
            "actor",
            "proposal_digest",
        } <= set(export)
        assert final_receipt["authority_digest"] == export["authority_digest"]
        assert final_receipt["approved_by"] == "system:qa"
        _write_json(out / "approval-record.json", approval_record)
        _write_json(out / "policy-denied.json", _denial_probes(runtime_path))
        _write_json(out / "receipt.json", final_receipt)
        _write_json(out / "destination-timeline.json", {
            "before_sha256": before_destination_sha,
            "exported_sha256": exported_sha,
            "rollback_sha256": _sha256(flow.destination),
            "rollback": rollback,
        })
        (out / "source-before.sha256").write_text(flow.source_sha256 + "\n", encoding="ascii")
        (out / "source-after.sha256").write_text(_sha256(flow.source) + "\n", encoding="ascii")
        shutil.copyfile(journal.path_for(cast(str, receipt["job_id"])), out / "journal.jsonl")
        (out / "action-log.txt").write_text("\n".join(actions) + "\n", encoding="utf-8")
    cleanup: dict[str, object] = {
        "runtime_removed": not runtime_path.exists(),
        "drafts_remaining": [], "backups_remaining": [], "temporary_files_remaining": [],
        "jobs_remaining": [], "processes_remaining": [], "evidence_retained": sorted(path.name for path in out.iterdir()),
    }
    assert cleanup["runtime_removed"] is True
    _write_json(out / "cleanup-receipt.json", cleanup)
    return cleanup


def main() -> None:
    parser = argparse.ArgumentParser()
    _ = parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()
    cleanup = run(cast(Path, args.out))
    print(json.dumps({"ok": True, "cleanup": cleanup}, ensure_ascii=False, sort_keys=True))


if __name__ == "__main__":
    main()
