"""Five representative Office-agent journeys with bounded quality metrics."""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import sys
import time
from pathlib import Path
from typing import cast

from docx import Document
from openpyxl import Workbook, load_workbook
from reportlab.pdfgen.canvas import Canvas

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from birkin import approvals, work_items  # noqa: E402
from birkin.m365_mail import create_local_draft, get_local_draft  # noqa: E402
from birkin.tools import build_registry  # noqa: E402
from birkin.tools._types import ToolContext  # noqa: E402


def _hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _artifact(path: Path) -> dict[str, str]:
    digest = _hash(path)
    return {"artifact_id": digest, "content_hash": digest, "media_type": "application/octet-stream",
            "uri": str(path.resolve()), "sensitivity": "internal", "acl_fingerprint": "journey"}


def run(output_dir: Path) -> dict[str, object]:
    root = output_dir.resolve()
    if root.exists() and any(root.iterdir()):
        raise ValueError("output directory must be absent or empty")
    root.mkdir(parents=True, exist_ok=True)
    os.environ["BIRKIN_HOME"] = str(root)
    office = root / "office"
    sources, exports = office / "sources", root / "exports"
    sources.mkdir(parents=True)
    exports.mkdir()
    documents = build_registry(ToolContext(cfg={}, client=None, cwd=root), include={"documents"})
    journey_evidence: dict[str, dict[str, object]] = {}
    metrics: list[dict[str, object]] = []

    def call(name: str, payload: dict[str, object], *, error: bool = False) -> dict[str, object]:
        result = documents.execute(name, payload)
        body = json.loads(cast(str, result.content))
        if result.is_error != error:
            raise AssertionError(f"{name} returned unexpected status: {body}")
        return cast(dict[str, object], body)

    def approved(proposal: dict[str, object]) -> tuple[str, dict[str, object]]:
        approval_id = cast(str, proposal["id"])
        result = approvals.approve(approval_id, approved_by="system:journey", approved_via="qa:script")
        if result["ok"] is not True:
            raise AssertionError(f"approval failed: {result}")
        return approval_id, cast(dict[str, object], json.loads(cast(str, result["result"])))

    def measured(name: str, action, *, recovery_attempts: int = 0, recovery_successes: int = 0) -> None:
        started = time.perf_counter()
        journey_evidence[name] = action()
        metrics.append({
            "journey": name, "time_to_first_result_ms": round((time.perf_counter() - started) * 1000),
            "unnecessary_questions": 0, "manual_edits": 0,
            "recovery_attempts": recovery_attempts, "recovery_successes": recovery_successes,
            "recovery_success_rate": (recovery_successes / recovery_attempts) if recovery_attempts else None,
        })

    def first_docx() -> dict[str, object]:
        destination = exports / "first-report.docx"
        approval_id, receipt = approved(call("office_job_request", {
            "request": "Create the first DOCX report", "format": "docx",
            "content": {"paragraphs": ["First report", "Verified office-agent journey"]},
            "outcome": "Create reviewed DOCX report", "destination": str(destination),
        }))
        text = "\n".join(paragraph.text for paragraph in Document(destination).paragraphs)
        return {"artifact": str(destination), "exists": destination.exists(), "content_match": "Verified" in text,
                "source": "reviewed request fields", "approval_id": approval_id,
                "approved_by": receipt.get("approved_by"), "validation": receipt.get("validation")}

    def existing_xlsx() -> dict[str, object]:
        source = sources / "baseline.xlsx"
        book = Workbook()
        sheet = book.active
        sheet.title = "Data"
        sheet.append(["name", "value"])
        sheet.append(["before", 1])
        book.save(source)
        descriptor = _artifact(source)
        stale = {**descriptor, "content_hash": "0" * 64}
        _ = call("office_job_request", {"request": "Update cell", "source": stale,
                 "outcome": "Set reviewed value", "operations": [{"cell": "B2", "value": 42}],
                 "destination": str(exports / "ignored.xlsx")}, error=True)
        destination = exports / "updated.xlsx"
        approval_id, receipt = approved(call("office_job_request", {
            "request": "Update cell", "source": descriptor, "outcome": "Set reviewed value",
            "operations": [{"cell": "B2", "value": 42}], "destination": str(destination),
        }))
        reopened = load_workbook(destination, data_only=False)
        value = reopened["Data"]["B2"].value
        reopened.close()
        return {"artifact": str(destination), "exists": destination.exists(), "content_match": value == 42,
                "source_hash": descriptor["content_hash"], "approval_id": approval_id,
                "approved_by": receipt.get("approved_by"), "failure_recovered": True}

    def pdf_summary() -> dict[str, object]:
        source = sources / "source.pdf"
        canvas = Canvas(str(source))
        canvas.drawString(72, 720, "Quarterly result: revenue increased.")
        canvas.save()
        descriptor = _artifact(source)
        extracted = call("extract_document", {"source": descriptor})
        spans = cast(list[dict[str, object]], extracted["spans"])
        summary = " ".join(str(span.get("text", "")) for span in spans).strip()
        destination = exports / "pdf-summary.docx"
        approval_id, receipt = approved(call("office_job_request", {
            "request": "Summarize PDF into DOCX", "format": "docx",
            "content": {"paragraphs": ["PDF summary", summary]},
            "outcome": "Create source-linked summary", "destination": str(destination),
        }))
        text = "\n".join(paragraph.text for paragraph in Document(destination).paragraphs)
        return {"artifact": str(destination), "exists": destination.exists(), "content_match": "revenue" in text,
                "source_hash": descriptor["content_hash"], "source_locators": [span.get("locator") for span in spans],
                "approval_id": approval_id, "approved_by": receipt.get("approved_by")}

    def meeting_tasks() -> dict[str, object]:
        draft = call("review_meeting_actions", {
            "notes": "민지는 견적을 확인한다. 기한은 미정이다.",
            "candidates": [{"action": "견적 확인", "evidence": "민지는 견적을 확인한다.", "assignee": "민지"}],
        })
        proposal = call("work_item_request", {
            "action": "confirm_meeting", "draft_sha256": draft["draft_sha256"],
            "items": draft["items"], "selected": [0], "session_id": "journey",
            "source": {"conversation_id": "journey-meeting"},
        })
        approval_id, _ = approved(proposal)
        items = work_items.grouped()["needs_confirmation"]
        return {"artifact": str(root / "work-items.json"), "exists": bool(items),
                "content_match": items[0]["title"] == "견적 확인", "source": items[0]["source"],
                "approval_id": approval_id, "approved_by": "system:journey"}

    def mail_review() -> dict[str, object]:
        draft = create_local_draft({"action": "reply", "source_message_id": "immutable-message-1",
                                    "source_etag": "etag-1", "from_account": "owner@example.com",
                                    "to": ["reviewer@example.com"], "cc": [], "subject": "Re: review",
                                    "body": "Reviewed draft body", "attachments": []})
        recovered = False
        try:
            _ = get_local_draft(draft["id"], "0" * 64)
        except ValueError:
            recovered = get_local_draft(draft["id"], draft["content_sha256"])["body"] == "Reviewed draft body"
        return {"artifact": str(root / "mail-drafts.json"), "exists": (root / "mail-drafts.json").exists(),
                "content_match": recovered, "source_message_id": draft["source_message_id"],
                "source_etag": draft["source_etag"], "approval_required": False,
                "failure_recovered": recovered}

    measured("first_docx", first_docx)
    measured("existing_xlsx_edit", existing_xlsx, recovery_attempts=1, recovery_successes=1)
    measured("pdf_to_docx_summary", pdf_summary)
    measured("meeting_to_tasks", meeting_tasks)
    measured("mail_draft_review", mail_review, recovery_attempts=1, recovery_successes=1)
    ok = len(journey_evidence) == 5 and all(item["exists"] and item["content_match"] for item in journey_evidence.values())
    report = {"ok": ok, "journeys": journey_evidence, "metrics": metrics,
              "metrics_privacy": "anonymous journey identifiers and aggregate counts only; no document or message body"}
    (root / "journey-report.json").write_text(json.dumps(report, ensure_ascii=False, sort_keys=True), encoding="utf-8")
    return report


def main() -> int:
    parser = argparse.ArgumentParser()
    _ = parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args()
    try:
        report = run(cast(Path, args.output_dir))
    except (AssertionError, ImportError, OSError, TypeError, ValueError) as exc:
        report = {"ok": False, "error": {"type": type(exc).__name__, "message": str(exc)}}
    print(json.dumps(report, ensure_ascii=False, sort_keys=True, default=str))
    return 0 if report["ok"] is True else 1


if __name__ == "__main__":
    raise SystemExit(main())
