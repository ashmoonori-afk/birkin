from __future__ import annotations

import hashlib
import zipfile
from pathlib import Path
from typing import cast

import pytest
from docx import Document

from birkin.office.adapters.docx import DocxAdapter
from birkin.office.adapters.docx_nodes import DocxLocator, DocxNode
from birkin.office.adapters.docx_types import DocxInspection
from birkin.office.errors import DocumentError, DocumentErrorCode


def _zip(path: Path, parts: dict[str, bytes]) -> Path:
    with zipfile.ZipFile(path, "w") as archive:
        for name, payload in parts.items():
            info = zipfile.ZipInfo(name, (2025, 2, 4, 6, 8, 10))
            info.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(info, payload)
    return path


def _parts(*, duplicate: bool = False, malformed: bool = False) -> dict[str, bytes]:
    duplicate_bookmark = (
        b'<w:bookmarkStart w:id="7" w:name="again"/><w:bookmarkEnd w:id="7"/>'
        if duplicate
        else b""
    )
    malformed_bookmark = b'<w:bookmarkStart w:name="bad"/>' if malformed else b""
    document = b"""<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>
<w:p w:rsidR="A"><w:pPr><w:pStyle w:val="Body"/></w:pPr><w:r><w:rPr><w:b/></w:rPr><w:t>Frag</w:t></w:r><w:r><w:rPr><w:i/></w:rPr><w:t>mented</w:t></w:r></w:p>
<w:p><w:bookmarkStart w:id="7" w:name="mark"/><w:r><w:t>Book</w:t></w:r><w:bookmarkEnd w:id="7"/></w:p>
<w:p><w:commentRangeStart w:id="1"/><w:r><w:t>Commented</w:t></w:r><w:commentRangeEnd w:id="1"/></w:p>
<w:p><w:fldSimple w:instr="simple"><w:r><w:t>SIMPLE</w:t></w:r></w:fldSimple></w:p>
<w:p><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText> complex </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>RESULT</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>
<w:sdt><w:sdtPr><w:tag w:val="first"/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>ONE</w:t></w:r></w:p></w:sdtContent></w:sdt>
<w:sdt><w:sdtPr><w:tag w:val="second"/></w:sdtPr><w:sdtContent><w:p><w:r><w:t>TWO</w:t></w:r></w:p></w:sdtContent></w:sdt>
<w:p><w:ins w:id="10" w:author="A" w:date="2025-01-01T00:00:00Z"><w:r><w:t>Inserted</w:t></w:r></w:ins><w:del w:id="11" w:author="A" w:date="2025-01-01T00:00:00Z"><w:r><w:delText>Deleted</w:delText></w:r></w:del><w:moveFrom w:id="12"><w:r><w:t>Moved</w:t></w:r></w:moveFrom></w:p>
<w:tbl><w:tr><w:tc><w:p><w:r><w:t>Outer</w:t></w:r></w:p><w:tbl><w:tr><w:tc><w:p><w:r><w:t>Nested</w:t></w:r></w:p></w:tc></w:tr></w:tbl></w:tc></w:tr></w:tbl>
""" + duplicate_bookmark + malformed_bookmark + b"<w:sectPr/></w:body></w:document>"
    return {
        "[Content_Types].xml": b"<Types/>",
        "word/document.xml": document,
        "word/header1.xml": b'<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:p><w:r><w:t>Header</w:t></w:r></w:p></w:hdr>',
        "word/footer1.xml": b'<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:p><w:r><w:t>Footer</w:t></w:r></w:p></w:ftr>',
        "word/footnotes.xml": b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:footnote w:id="-1"><w:p><w:r><w:t>Separator</w:t></w:r></w:p></w:footnote><w:footnote w:id="2"><w:p><w:r><w:t>Footnote</w:t></w:r></w:p></w:footnote></w:footnotes>',
        "word/endnotes.xml": b'<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:endnote w:id="2"><w:p><w:r><w:t>Endnote</w:t></w:r></w:p></w:endnote></w:endnotes>',
        "word/comments.xml": b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:comment w:id="1"><w:p><w:r><w:t>Comment body</w:t></w:r></w:p></w:comment></w:comments>',
        "word/styles.xml": b'<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:style w:styleId="Body"><w:name w:val="Body"/></w:style></w:styles>',
        "custom/unknown.xml": b"<unknown><sentinel/></unknown>",
        "custom/opaque.bin": b"opaque\x00sentinel",
    }


def _package(path: Path, **options: bool) -> Path:
    return _zip(path, _parts(**options))


def _hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        return {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
        }


def _paragraph(info: DocxInspection, story: str, text: str) -> DocxNode:
    paragraphs = info["paragraphs"]
    return next(node for node in paragraphs if node["story"] == story and node["text"] == text)


def test_inspection_has_stable_typed_locators_for_all_story_structures(tmp_path: Path) -> None:
    source = _package(tmp_path / "all.docx")
    adapter = DocxAdapter()
    first, second = adapter.inspect(source), adapter.inspect(source)
    assert first == second
    assert first["sections"] == 1
    assert first["headers"] == ["word/header1.xml"]
    assert first["footers"] == ["word/footer1.xml"]
    paragraphs = first["paragraphs"]
    assert {item["story"] for item in paragraphs} == {
        "body", "header", "footer", "footnote", "endnote", "comment"
    }
    assert _paragraph(first, "body", "Fragmented")["locator"]["kind"] == "paragraph"
    tables = first["tables"]
    assert sorted(table["table_depth"] for table in tables) == [0, 1]
    structures = cast("list[dict[str, object]]", first["structures"])
    assert {item["type"] for item in structures} >= {
        "bookmark", "comment_range", "simple", "complex", "content_control",
        "ins", "del", "moveFrom",
    }


@pytest.mark.parametrize(
    ("story", "old", "new", "part"),
    [
        ("body", "Fragmented", "Body changed", "word/document.xml"),
        ("header", "Header", "Head changed", "word/header1.xml"),
        ("footer", "Footer", "Foot changed", "word/footer1.xml"),
        ("footnote", "Footnote", "Note changed", "word/footnotes.xml"),
        ("endnote", "Endnote", "End changed", "word/endnotes.xml"),
        ("comment", "Comment body", "Reply changed", "word/comments.xml"),
    ],
)
def test_read_and_bounded_story_edit_change_only_the_target_part(
    tmp_path: Path, story: str, old: str, new: str, part: str
) -> None:
    source, output = _package(tmp_path / "source.docx"), tmp_path / "output.docx"
    adapter = DocxAdapter()
    node = _paragraph(adapter.inspect(source), story, old)
    locator = node["locator"]
    assert adapter.read(source, locator)["text"] == old
    before = _hashes(source)
    evidence = adapter.patch_text(source, output, locator, new, expected_text=old)
    after = _hashes(output)
    assert evidence["source_part"] == part
    assert before[part] != after[part]
    assert {key: value for key, value in before.items() if key != part} == {
        key: value for key, value in after.items() if key != part
    }


def test_single_run_edit_preserves_sibling_run_and_style_wrappers(tmp_path: Path) -> None:
    source, output = _package(tmp_path / "source.docx"), tmp_path / "run.docx"
    adapter, info = DocxAdapter(), DocxAdapter().inspect(source)
    target = next(node for node in info["runs"] if node["story"] == "body" and node["text"] == "Frag")
    assert adapter.read_node(source, target["locator"])["text"] == "Frag"
    _ = adapter.edit_text(source, output, target["locator"], "Piece", expected_text="Frag")
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml")
    assert b">Piece</w:t>" in xml and b">mented</w:t>" in xml
    assert b"<w:rPr><w:b/></w:rPr>" in xml and b"<w:rPr><w:i/></w:rPr>" in xml


@pytest.mark.parametrize("ancestor", ["ins", "tbl"])
def test_field_edit_refuses_namespaced_alias_ancestor(
    tmp_path: Path, ancestor: str
) -> None:
    parts = _parts()
    target = (
        b'<w:sdt><w:sdtPr><w:tag w:val="first"/></w:sdtPr>'
        b'<w:sdtContent><w:p><w:r><w:t>ONE</w:t></w:r></w:p></w:sdtContent></w:sdt>'
    )
    namespace = b"http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    wrapped = b'<guard:' + ancestor.encode() + b' xmlns:guard="' + namespace + b'">' + target + b'</guard:' + ancestor.encode() + b'>'
    parts["word/document.xml"] = parts["word/document.xml"].replace(target, wrapped)
    source = _zip(tmp_path / f"aliased-{ancestor}.docx", parts)
    output = tmp_path / "refused.docx"

    with pytest.raises(DocumentError) as caught:
        _ = DocxAdapter().patch_field(source, output, "first", "unsafe")

    assert caught.value.code is DocumentErrorCode.UNSUPPORTED_EDIT
    assert not output.exists()


@pytest.mark.parametrize("ancestor", ["ins", "tbl"])
def test_field_edit_refuses_alias_guard_with_entity_encoded_word_namespace(
    tmp_path: Path, ancestor: str
) -> None:
    base = tmp_path / "base.docx"
    document = Document()
    _ = document.add_paragraph("reopen sentinel")
    document.save(str(base))
    with zipfile.ZipFile(base) as archive:
        parts = {name: archive.read(name) for name in archive.namelist()}
    target = (
        b'<w:sdt><w:sdtPr><w:tag w:val="guarded"/></w:sdtPr>'
        b'<w:sdtContent><w:r><w:t>LOCKED</w:t></w:r></w:sdtContent></w:sdt>'
    )
    namespace = b"http&#x3A;//schemas.openxmlformats.org/wordprocessingml/2006/main"
    if ancestor == "ins":
        guarded = (
            b'<w:p><alias:ins xmlns:alias="' + namespace + b'" w:id="91">'
            + target + b"</alias:ins></w:p>"
        )
    else:
        guarded = (
            b'<alias:tbl xmlns:alias="' + namespace + b'"><w:tr><w:tc><w:p>'
            + target + b"</w:p></w:tc></w:tr></alias:tbl>"
        )
    parts["word/document.xml"] = parts["word/document.xml"].replace(
        b"<w:sectPr", guarded + b"<w:sectPr", 1
    )
    source = _zip(tmp_path / f"encoded-{ancestor}.docx", parts)
    assert Document(str(source)).paragraphs[0].text == "reopen sentinel"

    output = tmp_path / "refused.docx"
    with pytest.raises(DocumentError) as caught:
        _ = DocxAdapter().patch_field(source, output, "guarded", "unsafe")

    assert caught.value.code is DocumentErrorCode.UNSUPPORTED_EDIT
    assert not output.exists()


@pytest.mark.parametrize("ancestor", ["ins", "tbl"])
def test_text_edit_refuses_alias_guard_with_entity_encoded_word_namespace(
    tmp_path: Path, ancestor: str
) -> None:
    base = tmp_path / "base-edit.docx"
    document = Document()
    _ = document.add_paragraph("reopen sentinel")
    document.save(str(base))
    with zipfile.ZipFile(base) as archive:
        parts = {name: archive.read(name) for name in archive.namelist()}
    run = b"<w:r><w:t>LOCKED EDIT</w:t></w:r>"
    namespace = b"http&#58;//schemas.openxmlformats.org/wordprocessingml/2006/main"
    if ancestor == "ins":
        guarded = (
            b'<w:p><alias:ins xmlns:alias="' + namespace + b'" w:id="92">'
            + run + b"</alias:ins></w:p>"
        )
    else:
        guarded = (
            b'<alias:tbl xmlns:alias="' + namespace + b'"><w:tr><w:tc><w:p>'
            + run + b"</w:p></w:tc></w:tr></alias:tbl>"
        )
    parts["word/document.xml"] = parts["word/document.xml"].replace(
        b"<w:sectPr", guarded + b"<w:sectPr", 1
    )
    source = _zip(tmp_path / f"encoded-edit-{ancestor}.docx", parts)
    _ = Document(str(source))
    target = next(
        node for node in DocxAdapter().inspect(source)["runs"]
        if node["text"] == "LOCKED EDIT"
    )

    output = tmp_path / "refused-edit.docx"
    with pytest.raises(DocumentError) as caught:
        _ = DocxAdapter().patch_text(source, output, target["locator"], "unsafe")

    assert caught.value.code is DocumentErrorCode.UNSUPPORTED_EDIT
    assert not output.exists()


def test_entity_aliased_duplicate_content_control_is_ambiguous_before_splice(
    tmp_path: Path,
) -> None:
    base = tmp_path / "base-duplicate.docx"
    document = Document()
    _ = document.add_paragraph("reopen sentinel")
    document.save(str(base))
    with zipfile.ZipFile(base) as archive:
        parts = {name: archive.read(name) for name in archive.namelist()}
    canonical = (
        b'<w:sdt><w:sdtPr><w:tag w:val="duplicate"/></w:sdtPr>'
        b'<w:sdtContent><w:r><w:t>ONE</w:t></w:r></w:sdtContent></w:sdt>'
    )
    namespace = b"http&#x3A;//schemas.openxmlformats.org/wordprocessingml/2006/main"
    aliased = (
        b'<review:sdt xmlns:review="' + namespace + b'">'
        b'<review:sdtPr><review:tag review:val="duplicate"/></review:sdtPr>'
        b'<review:sdtContent><review:r><review:t>TWO</review:t></review:r>'
        b'</review:sdtContent></review:sdt>'
    )
    parts["word/document.xml"] = parts["word/document.xml"].replace(
        b"<w:sectPr", canonical + aliased + b"<w:sectPr", 1
    )
    source = _zip(tmp_path / "entity-aliased-duplicate.docx", parts)
    _ = Document(str(source))
    before = source.read_bytes()
    output = tmp_path / "ambiguous.docx"

    with pytest.raises(DocumentError) as caught:
        _ = DocxAdapter().patch_field(
            source, output, "duplicate", "changed", expected_text="ONE"
        )

    assert caught.value.code is DocumentErrorCode.AMBIGUOUS_LOCATOR
    assert caught.value.stage == "locate"
    assert source.read_bytes() == before
    assert not output.exists()
    _ = Document(str(source))


def test_template_merge_is_hash_bound_explicit_and_preserves_unrelated_parts(tmp_path: Path) -> None:
    source, output = _package(tmp_path / "template.docx"), tmp_path / "merged.docx"
    digest = hashlib.sha256(source.read_bytes()).hexdigest()
    before = _hashes(source)
    result = DocxAdapter().merge_template(
        source,
        output,
        {"first": {"value": "Ada", "expected_text": "ONE"}, "second": "Lovelace"},
        expected_template_sha256=digest,
    )
    after = _hashes(output)
    assert result["template_sha256"] == digest
    assert result["changed_parts"] == ["word/document.xml"]
    assert before["custom/opaque.bin"] == after["custom/opaque.bin"]
    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml")
    assert b">Ada</w:t>" in xml and b">Lovelace</w:t>" in xml
    with pytest.raises(DocumentError) as stale:
        _ = DocxAdapter().merge_template(
            source, tmp_path / "stale.docx", {"first": "x"}, expected_template_sha256="0" * 64
        )
    assert stale.value.code is DocumentErrorCode.SOURCE_CHANGED


def test_table_separator_field_revision_range_style_and_stale_cases_refuse(tmp_path: Path) -> None:
    source, output = _package(tmp_path / "source.docx"), tmp_path / "output.docx"
    adapter, info = DocxAdapter(), DocxAdapter().inspect(source)
    tables = info["tables"]
    separator = _paragraph(info, "footnote", "Separator")
    revision = _paragraph(info, "body", "InsertedDeletedMoved")
    field = _paragraph(info, "body", "SIMPLE")
    for locator in (
        tables[0]["locator"], separator["locator"], revision["locator"], field["locator"]
    ):
        with pytest.raises(DocumentError) as caught:
            _ = adapter.patch_text(source, output, locator, "unsafe")
        assert caught.value.code is DocumentErrorCode.UNSUPPORTED_EDIT
        assert not output.exists()
    with pytest.raises(DocumentError) as table_edit:
        adapter.patch_table(source, output, tables[0]["locator"], "unsafe")
    assert table_edit.value.details["operation"] == "edit_table"


def test_typed_refusals_and_duplicate_malformed_inventory(tmp_path: Path) -> None:
    source = _package(tmp_path / "bad.docx", duplicate=True, malformed=True)
    info = DocxAdapter().inspect(source)
    bookmarks = info["bookmarks"]
    assert any(item["state"] == "malformed" and "duplicate_id" in item["reasons"] for item in bookmarks)
    assert any(item["state"] == "malformed" and "missing_id" in item["reasons"] for item in bookmarks)
    output = tmp_path / "out.docx"
    for action in ("create", "accept", "reject"):
        with pytest.raises(DocumentError) as caught:
            DocxAdapter().patch_tracked_change(source, output, "10", action=action)
        assert caught.value.code is DocumentErrorCode.UNSUPPORTED_EDIT
        assert caught.value.details["operation"] == f"tracked_change_{action}"
    with pytest.raises(DocumentError) as invalid:
        DocxAdapter().patch_tracked_change(source, output, "10", action="maybe")
    assert invalid.value.code is DocumentErrorCode.INVALID_INPUT


def test_locator_and_expected_text_preconditions_are_enforced(tmp_path: Path) -> None:
    first = _package(tmp_path / "first.docx")
    second_parts = _parts()
    second_parts["custom/opaque.bin"] = b"different"
    second = _zip(tmp_path / "second.docx", second_parts)
    locator: DocxLocator = _paragraph(DocxAdapter().inspect(first), "body", "Fragmented")["locator"]
    with pytest.raises(DocumentError) as source_changed:
        _ = DocxAdapter().read(second, locator)
    assert source_changed.value.code is DocumentErrorCode.SOURCE_CHANGED
    with pytest.raises(DocumentError) as stale_text:
        _ = DocxAdapter().patch_text(first, tmp_path / "out.docx", locator, "new", expected_text="stale")
    assert stale_text.value.code is DocumentErrorCode.PRECONDITION_FAILED
