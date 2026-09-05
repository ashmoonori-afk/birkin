from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import cast

import pytest
from docx import Document

from birkin import approvals, store
from birkin.office import create_execution
from birkin.office.create_journal import CreationJobJournal
from birkin.office.errors import DocumentError, DocumentErrorCode
from birkin.office.service import DocumentService
from birkin.tools import build_registry
from birkin.tools._types import ToolContext
from birkin.workspace.approval_projection import approval_item


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _queue_creation(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> tuple[dict[str, object], dict[str, object], Path, Path]:
    home = tmp_path / "home"
    caller = tmp_path / "caller"
    office_home = home / "office"
    office_home.mkdir(parents=True)
    caller.mkdir()
    monkeypatch.setenv("BIRKIN_HOME", str(home))
    destination = caller / "quarterly-report.docx"
    registry = build_registry(
        ToolContext(
            cfg={},
            client=None,
            cwd=caller,
            record_source="user:office-create-e2e",
        ),
        include={"documents"},
    )

    proposed = registry.execute(
        "office_job_request",
        {
            "request": "2026년 3분기 보고서를 새 DOCX로 만들어 주세요.",
            "format": "docx",
            "content": {
                "paragraphs": [
                    "2026년 3분기 보고서",
                    "매출은 전분기 대비 18% 증가했습니다.",
                ],
            },
            "outcome": "새 분기 보고서 작성",
            "destination": str(destination),
        },
    )

    assert isinstance(proposed.content, str)
    body = cast("dict[str, object]", json.loads(proposed.content))
    assert not proposed.is_error, body
    approval_id = cast(str, body["id"])
    record = store.get_pending(approval_id)
    assert record is not None
    return body, record, office_home, destination


def test_creation_request_queues_approval_without_writing_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Given: a request for a brand-new report in an allowlisted directory.
    body, record, office_home, destination = _queue_creation(
        tmp_path,
        monkeypatch,
    )

    # When: the queued approval is inspected before any human decision.
    payload = cast("dict[str, object]", record["payload"])
    approval = cast("dict[str, object]", body["approval"])

    # Then: creation authority is bound without creating a draft or destination.
    assert record["category"] == "office_create"
    assert payload == approval
    assert payload["format"] == "docx"
    assert payload["destination"] == str(destination)
    assert payload["content_sha256"]
    assert payload["creation_digest"]
    assert payload["authority_digest"]
    assert approval_item(record)["sealed"] is True
    journal = CreationJobJournal(office_home).latest(cast(str, payload["job_id"]))
    assert journal["state"] == "approval_requested"
    assert journal["approval"] == payload
    assert not destination.exists()
    assert not list((office_home / "artifacts" / "drafts").iterdir())


def test_source_xlsx_and_target_docx_queue_one_creation_approval(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    home = tmp_path / "home"
    caller = tmp_path / "caller"
    (home / "office").mkdir(parents=True)
    caller.mkdir()
    monkeypatch.setenv("BIRKIN_HOME", str(home))
    destination = caller / "comparison-report.docx"
    registry = build_registry(
        ToolContext(
            cfg={},
            client=None,
            cwd=caller,
            record_source="user:mixed-format-clarification",
        ),
        include={"documents"},
    )

    result = registry.execute(
        "office_job_request",
        {
            "request": "엑셀 비교해서 워드 보고서로",
            "format": "docx",
            "content": {"paragraphs": ["비교 보고서"]},
            "outcome": "비교 보고서 작성",
            "destination": str(destination),
        },
    )

    assert result.is_error is False
    body = cast("dict[str, object]", json.loads(cast(str, result.content)))
    assert body["category"] == "office_create"
    assert not destination.exists()


def test_output_exists_queues_one_click_creation_overwrite_approval(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    body, original_record, _, destination = _queue_creation(
        tmp_path,
        monkeypatch,
    )
    existing = b"existing caller file"
    destination.write_bytes(existing)

    first = approvals.approve(
        cast(str, body["id"]),
        approved_by="human:first-reviewer",
        approved_via="test:first-approval",
    )

    assert first["ok"] is False
    assert first["error"] == "기존 파일을 덮어쓸까요?"
    follow_up_id = cast(str, first["follow_up_approval_id"])
    assert destination.read_bytes() == existing
    resolved = store.get_pending(cast(str, body["id"]))
    assert resolved is not None
    assert resolved["status"] == "error"
    assert resolved["failure_code"] == "OUTPUT_EXISTS"
    assert resolved["follow_up_approval_id"] == follow_up_id
    follow_up = store.get_pending(follow_up_id)
    assert follow_up is not None
    assert follow_up["title"] == "기존 파일을 덮어쓸까요?"
    assert follow_up["retry_of_approval_id"] == body["id"]
    assert follow_up["overwrite_retry"] is True
    payload = cast("dict[str, object]", follow_up["payload"])
    original_payload = cast("dict[str, object]", original_record["payload"])
    assert payload["overwrite_approved"] is True
    assert payload["content"] == original_payload["content"]
    assert payload["destination"] == original_payload["destination"]
    assert payload["job_id"] != original_payload["job_id"]

    second = approvals.approve(
        follow_up_id,
        approved_by="human:overwrite-reviewer",
        approved_via="test:overwrite-approval",
    )

    assert second["ok"] is True
    assert [paragraph.text for paragraph in Document(str(destination)).paragraphs] == [
        "2026년 3분기 보고서",
        "매출은 전분기 대비 18% 증가했습니다.",
    ]


def test_approved_creation_writes_real_docx_and_hash_receipt(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Given: a bound creation proposal with no pre-approval output.
    body, _, office_home, destination = _queue_creation(
        tmp_path,
        monkeypatch,
    )
    assert not destination.exists()

    # When: the standard approval queue executes the creation proposal.
    result = approvals.approve(
        cast(str, body["id"]),
        approved_by="human:test-reviewer",
        approved_via="test:office-create",
    )

    # Then: the exported DOCX reopens with the approved content and hash proof.
    assert result["ok"] is True, result
    receipt = cast("dict[str, object]", json.loads(cast(str, result["result"])))
    exported = cast("dict[str, object]", receipt["export"])
    assert receipt["state"] == "exported"
    assert destination.is_file()
    assert [paragraph.text for paragraph in Document(str(destination)).paragraphs] == [
        "2026년 3분기 보고서",
        "매출은 전분기 대비 18% 증가했습니다.",
    ]
    assert exported["output_sha256"] == _sha256(destination)
    assert exported["path"] == str(destination)
    journal = CreationJobJournal(office_home).latest(cast(str, receipt["job_id"]))
    assert journal["state"] == "exported"
    assert journal["export"] == exported


def test_creation_execution_is_denied_outside_the_approval_queue(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Given: a creation proposal that is still pending human approval.
    _, record, _, destination = _queue_creation(tmp_path, monkeypatch)
    payload = cast("dict[str, object]", record["payload"])

    # When: execution is attempted without the approval queue transition.
    with pytest.raises(DocumentError) as caught:
        _ = approvals.execute_action("office_create", payload)

    # Then: the attempt fails closed without creating either output.
    assert caught.value.code is DocumentErrorCode.PRECONDITION_FAILED
    assert not destination.exists()


def test_creation_content_tamper_is_denied_before_document_write(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Given: an executing approval whose content no longer matches its digest.
    body, record, office_home, destination = _queue_creation(
        tmp_path,
        monkeypatch,
    )
    tampered = cast(
        "dict[str, object]",
        json.loads(json.dumps(record["payload"])),
    )
    content = cast("dict[str, object]", tampered["content"])
    content["paragraphs"] = ["승인되지 않은 내용"]
    monkeypatch.setattr(
        create_execution.store,
        "get_pending",
        lambda _approval_id: {
            **record,
            "status": "executing",
            "approved_by": "human:test-reviewer",
            "approved_via": "test:tamper",
            "payload": tampered,
        },
    )

    # When: the tampered payload reaches the approved executor.
    with pytest.raises(DocumentError) as caught:
        _ = create_execution.execute_approved_office_creation(
            tampered,
            approval_id=cast(str, body["id"]),
        )

    # Then: digest verification stops execution before any DOCX write.
    assert caught.value.code is DocumentErrorCode.PRECONDITION_FAILED
    assert "content changed" in caught.value.message
    assert not destination.exists()
    assert not list((office_home / "artifacts" / "drafts").iterdir())


def test_creation_rejects_truncated_preexisting_draft(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Given: a managed draft with approved paragraphs plus hidden extra content.
    body, record, office_home, destination = _queue_creation(
        tmp_path,
        monkeypatch,
    )
    payload = cast("dict[str, object]", record["payload"])
    _ = DocumentService(office_home).create_document(
        format="docx",
        content={
            "paragraphs": [
                "2026년 3분기 보고서",
                "매출은 전분기 대비 18% 증가했습니다.",
                "승인되지 않은 추가 문단",
            ],
        },
        output_name=cast(str, payload["output_name"]),
    )
    monkeypatch.setattr(
        create_execution.store,
        "get_pending",
        lambda _approval_id: {
            **record,
            "status": "executing",
            "approved_by": "human:test-reviewer",
            "approved_via": "test:truncated-draft",
        },
    )

    # When: execution reuses and extracts the deterministic managed draft.
    with pytest.raises(DocumentError) as caught:
        _ = create_execution.execute_approved_office_creation(
            payload,
            approval_id=cast(str, body["id"]),
        )

    # Then: truncated semantic extraction cannot authorize an export.
    assert caught.value.code is DocumentErrorCode.SOURCE_CHANGED
    assert not destination.exists()


def test_approved_creation_rolls_back_through_second_approval(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Given: an approved creation with an authenticated durable receipt.
    body, _, office_home, destination = _queue_creation(
        tmp_path,
        monkeypatch,
    )
    created = approvals.approve(
        cast(str, body["id"]),
        approved_by="human:create-reviewer",
        approved_via="test:create",
    )
    creation_receipt = cast(
        "dict[str, object]",
        json.loads(cast(str, created["result"])),
    )
    registry = build_registry(
        ToolContext(
            cfg={},
            client=None,
            cwd=destination.parent,
            record_source="user:office-create-e2e",
        ),
        include={"documents"},
    )

    # When: rollback is separately requested and approved by a human.
    proposed = registry.execute(
        "office_rollback_request",
        {"job_id": creation_receipt["job_id"]},
    )
    rollback = cast(
        "dict[str, object]",
        json.loads(cast(str, proposed.content)),
    )
    assert not proposed.is_error, rollback
    assert destination.is_file()
    result = approvals.approve(
        cast(str, rollback["id"]),
        approved_by="human:rollback-reviewer",
        approved_via="test:create-rollback",
    )

    # Then: rollback removes the newly created destination and is durable.
    assert result["ok"] is True, result
    receipt = cast(
        "dict[str, object]",
        json.loads(cast(str, result["result"])),
    )
    assert receipt["restored"] is False
    assert receipt["destination_sha256"] is None
    assert not destination.exists()
    journal = CreationJobJournal(office_home).latest(
        cast(str, creation_receipt["job_id"])
    )
    assert journal["state"] == "rolled_back"
    assert journal["rollback"] == receipt
