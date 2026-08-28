from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path
from typing import cast

import pytest
from docx import Document

from birkin import approvals, config, store
from birkin.office.errors import DocumentError, DocumentErrorCode
from birkin.tools import build_registry
from birkin.tools._types import ToolContext


def _xlsx(path: Path) -> Path:
    parts = {
        "[Content_Types].xml": b'<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/></Types>',
        "xl/workbook.xml": b'<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets><sheet name="Revenue" sheetId="1" r:id="rId1"/></sheets></workbook>',
        "xl/_rels/workbook.xml.rels": b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/></Relationships>',
        "xl/worksheets/sheet1.xml": b'<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"><sheetData><row r="1"><c r="A1"><v>7</v></c></row></sheetData></worksheet>',
    }
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, payload in parts.items():
            archive.writestr(name, payload)
    return path


def _request(home: Path, destination: Path) -> tuple[dict[str, object], Path, str]:
    source_path = _xlsx(home / "source.xlsx")
    source_sha256 = hashlib.sha256(source_path.read_bytes()).hexdigest()
    return (
        {
            "request": "Update cell A1 in this Excel workbook",
            "source": {
                "artifact_id": source_sha256,
                "content_hash": source_sha256,
                "media_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "uri": str(source_path),
                "sensitivity": "internal",
                "acl_fingerprint": "local-contract",
            },
            "outcome": "Set Revenue A1 to 9",
            "operations": [{"cell": "A1", "value": 9}],
            "destination": str(destination),
        },
        source_path,
        source_sha256,
    )


def queue_office_job(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> tuple[dict[str, object], dict[str, object], Path, Path, str]:
    home = tmp_path / "home"
    export_root = tmp_path / "caller"
    office_home = home / "office"
    office_home.mkdir(parents=True)
    export_root.mkdir()
    monkeypatch.setenv("BIRKIN_HOME", str(home))
    destination = export_root / "approved.xlsx"
    request, source, source_sha256 = _request(office_home, destination)
    context = ToolContext(
        cfg={},
        client=None,
        cwd=export_root,
        record_source="user:local-contract",
    )

    result = build_registry(context, include={"documents"}).execute(
        "office_job_request", request
    )

    assert isinstance(result.content, str)
    body = cast("dict[str, object]", json.loads(result.content))
    assert not result.is_error, body
    approval_id = cast(str, body["id"])
    record = store.get_pending(approval_id)
    assert record is not None
    return body, record, source, destination, source_sha256


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_docx_paragraph_request_executes_through_registry_and_approval(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    # Given: a real DOCX and a natural-language bilingual paragraph request.
    home = tmp_path / "home"
    caller = tmp_path / "caller"
    office_home = home / "office"
    office_home.mkdir(parents=True)
    caller.mkdir()
    monkeypatch.setenv("BIRKIN_HOME", str(home))
    source = office_home / "source.docx"
    destination = caller / "approved.docx"
    document = Document()
    _ = document.add_paragraph("계약 원문 Original contract paragraph.")
    document.save(str(source))
    source_sha256 = _sha256(source)
    request = {
        "request": (
            "이 DOCX Word document의 계약 원문 Original contract paragraph를 "
            "승인된 계약문 Approved contract paragraph로 교체해 주세요."
        ),
        "source": {"content_hash": source_sha256, "uri": str(source)},
        "outcome": "Replace the known bilingual contract paragraph",
        "operations": [
            {
                "locator": {"format": "docx", "index": 1},
                "value": "승인된 계약문 Approved contract paragraph.",
            }
        ],
        "destination": str(destination),
    }
    registry = build_registry(
        ToolContext(
            cfg={}, client=None, cwd=caller, record_source="user:docx-e2e"
        ),
        include={"documents"},
    )

    # When: the registry queues and the standard approval queue executes the job.
    proposed = registry.execute("office_job_request", request)
    body = cast("dict[str, object]", json.loads(cast(str, proposed.content)))
    result = approvals.approve(
        cast(str, body["id"]),
        approved_by="human:test-reviewer",
        approved_via="test:office-coordinator",
    )

    # Then: the exported DOCX reopens with only the approved paragraph replacement.
    assert not proposed.is_error, body
    assert result["ok"] is True, result
    assert Document(str(destination)).paragraphs[0].text == (
        "승인된 계약문 Approved contract paragraph."
    )
    assert _sha256(source) == source_sha256


def test_request_queues_bound_approval_without_mutating_files(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    # Given: one source and a caller-owned allowlisted destination.
    body, record, source, destination, source_sha256 = queue_office_job(tmp_path, monkeypatch)

    # When: the canonical approval request is inspected before approval.
    payload = cast("dict[str, object]", record["payload"])
    approval = cast("dict[str, object]", body["approval"])

    # Then: every execution authority is bound and no mutation has happened.
    assert record["category"] == "office_job"
    assert payload == approval
    assert {
        "job_id",
        "proposal_digest",
        "authority_digest",
        "source_sha256",
        "destination",
        "allowlist_root",
        "proposer",
    } <= set(payload)
    assert payload["source_sha256"] == source_sha256
    assert payload["destination"] == str(destination)
    assert payload["allowlist_root"] == str(destination.parent)
    assert payload["proposer"] == "user:local-contract"
    assert isinstance(payload["authority_digest"], str)
    summaries = cast("list[dict[str, str]]", payload["semantic_summaries"])
    assert len(summaries) == 1
    assert summaries[0]["location"] == "A1"
    assert _sha256(source) == source_sha256
    assert not destination.exists()


def test_direct_resume_is_policy_denied_without_file_mutation(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    # Given: a durable Office job still waiting in the canonical queue.
    _, record, source, destination, source_sha256 = queue_office_job(tmp_path, monkeypatch)
    payload = cast("dict[str, object]", record["payload"])

    # When: execution is attempted outside the approved queue transition.
    with pytest.raises(DocumentError) as caught:
        _ = approvals.execute_action("office_job", payload)

    # Then: policy denial is typed and both filesystem identities are unchanged.
    assert caught.value.code is DocumentErrorCode.POLICY_DENIED
    assert _sha256(source) == source_sha256
    assert not destination.exists()


def test_approved_queue_executes_validates_materializes_and_exports(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    # Given: a bound Office proposal and unchanged source/destination hashes.
    body, _, source, destination, source_sha256 = queue_office_job(tmp_path, monkeypatch)

    # When: the canonical approval queue executes the exact proposal.
    result = approvals.approve(
        cast(str, body["id"]),
        approved_by="human:test-reviewer",
        approved_via="test:office-coordinator",
    )

    # Then: a validated internal artifact is exported with real hash proof.
    assert result["ok"] is True, result
    receipt = cast("dict[str, object]", json.loads(cast(str, result["result"])))
    publication = cast("dict[str, object]", receipt["publication"])
    exported = cast("dict[str, object]", receipt["export"])
    internal = Path(cast(str, publication["path"]))
    assert receipt["state"] == "exported"
    assert internal.is_file()
    assert destination.is_file()
    assert _sha256(source) == source_sha256
    assert _sha256(internal) == publication["sha256"]
    assert _sha256(destination) == exported["output_sha256"]
    assert exported["output_sha256"] == publication["sha256"]
    assert exported["actor"] == "user:local-contract"
    assert receipt["authority_digest"] == exported["authority_digest"]
    assert receipt["approved_by"] == "human:test-reviewer"


def test_digest_drift_is_denied_before_mutation(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    # Given: a queued record whose approved digest no longer matches its durable job.
    body, record, source, destination, source_sha256 = queue_office_job(tmp_path, monkeypatch)
    payload = cast("dict[str, object]", record["payload"])
    record["payload"] = {**payload, "proposal_digest": "0" * 64}
    config.pending_dir().joinpath(f"{body['id']}.json").write_text(
        json.dumps(record), encoding="utf-8"
    )

    # When: the drifted record is approved through the canonical queue.
    result = approvals.approve(
        cast(str, body["id"]),
        approved_by="human:test-reviewer",
        approved_via="test:office-coordinator",
    )

    # Then: execution is refused before any draft or destination is written.
    assert result["ok"] is False
    assert "POLICY_DENIED" in cast(str, result["error"])
    assert _sha256(source) == source_sha256
    assert not destination.exists()


def test_destination_drift_is_denied_before_mutation(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    # Given: a queued approval whose destination changed after review.
    body, record, source, destination, source_sha256 = queue_office_job(tmp_path, monkeypatch)
    payload = cast("dict[str, object]", record["payload"])
    drifted_destination = destination.with_name("drifted.docx")
    record["payload"] = {
        **payload,
        "destination": str(drifted_destination),
    }
    config.pending_dir().joinpath(f"{body['id']}.json").write_text(
        json.dumps(record), encoding="utf-8"
    )

    # When: the mutated approval record reaches the execution boundary.
    result = approvals.approve(
        cast(str, body["id"]),
        approved_by="human:test-reviewer",
        approved_via="test:office-coordinator",
    )

    # Then: authority verification rejects it before either destination changes.
    assert result["ok"] is False
    assert "POLICY_DENIED" in cast(str, result["error"])
    assert _sha256(source) == source_sha256
    assert not destination.exists()
    assert not drifted_destination.exists()
