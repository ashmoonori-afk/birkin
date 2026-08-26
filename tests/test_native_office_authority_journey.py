from __future__ import annotations

import hashlib
import zipfile
from pathlib import Path
from typing import cast

import pytest
from docx import Document
from openpyxl import load_workbook

from birkin import store
from birkin.workspace.contracts import ClientContext, WorkspaceCommand
from birkin.workspace.records import CommandReceipt
from birkin.workspace.runtime_adapter import RuntimeWorkspaceAdapter
from birkin.workspace.service import WorkspaceService


FIXTURES = (
    Path("windows/BirkinNativeApp/tests/Birkin.Native.App.Tests/Fixtures/Office")
)


def _service(
    tmp_path: Path,
) -> tuple[WorkspaceService, RuntimeWorkspaceAdapter]:
    service = WorkspaceService(
        root=tmp_path / "journal", session_id="office-journey", handlers={}
    )
    adapter = RuntimeWorkspaceAdapter(
        "office-journey", service.emit, workspace_root=tmp_path / "workspace"
    )
    service.set_handlers(adapter.handlers())
    return service, adapter


def _submit(
    service: WorkspaceService,
    command_id: str,
    command_type: str,
    payload: dict[str, object],
) -> tuple[CommandReceipt, dict[str, object]]:
    command = WorkspaceCommand(
        protocol_version=1,
        command_id=command_id,
        expected_cursor=service.snapshot().cursor,
        type=command_type,
        payload=payload,
        client_context=ClientContext(surface="windows", view_id="office"),
    )
    receipt = service.submit(command, actor_id="windows:office")
    result = receipt.transient_result
    assert receipt.state == "completed"
    assert isinstance(result, dict)
    return receipt, result


def _artifact(result: dict[str, object]) -> dict[str, object]:
    artifact = result.get("artifact")
    assert isinstance(artifact, dict)
    return cast(dict[str, object], artifact)


def test_file_import_registers_office_artifacts_by_copy_without_client_bytes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("BIRKIN_HOME", str(tmp_path / "home"))
    service, adapter = _service(tmp_path)
    source = tmp_path / "outside" / "baseline.xlsx"
    source.parent.mkdir()
    _ = source.write_bytes((FIXTURES / "baseline.xlsx").read_bytes())

    _receipt, imported = _submit(
        service, "import-baseline", "file.import", {"source_path": str(source)}
    )
    artifact = _artifact(imported)
    registered = Path(cast(str, artifact["uri"]))
    _ = source.write_bytes(b"changed after authoritative copy")

    assert set(imported) == {"reference", "artifact", "receipt"}
    assert registered.is_relative_to(adapter.surface_authority.office.service.home)
    assert registered.read_bytes() != source.read_bytes()
    assert hashlib.sha256(registered.read_bytes()).hexdigest() == artifact["content_hash"]
    assert str(source) not in str(imported)

    with pytest.raises(ValueError, match="canonical source_path"):
        _ = _submit(
            service,
            "import-traversal",
            "file.import",
            {"source_path": str(source), "destination": "../escape.xlsx"},
        )
    assert not (tmp_path / "escape.xlsx").exists()


def test_python_authority_runs_import_diff_approval_save_and_activity_receipt(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("BIRKIN_HOME", str(tmp_path / "home"))
    service, adapter = _service(tmp_path)
    changed_candidate = tmp_path / "changed-candidate.xlsx"
    workbook = load_workbook(FIXTURES / "candidate.xlsx")
    sheet = workbook["Comparison"]
    sheet["B1"] = 4800
    sheet["B2"] = "office-4800"
    workbook.save(changed_candidate)
    sources = {
        "baseline.xlsx": FIXTURES / "baseline.xlsx",
        "candidate.xlsx": changed_candidate,
        "report-template.docx": FIXTURES / "report-template.docx",
    }
    imported: dict[str, dict[str, object]] = {}
    for name, source in sources.items():
        _receipt, result = _submit(
            service,
            f"import-{name.split('.')[0]}",
            "file.import",
            {"source_path": str(source.resolve())},
        )
        imported[name] = _artifact(result)

    compare_receipt, compared = _submit(
        service,
        "compare-spreadsheets",
        "office.compare",
        {
            "left_artifact_id": imported["baseline.xlsx"]["artifact_id"],
            "right_artifact_id": imported["candidate.xlsx"]["artifact_id"],
        },
    )
    diff = cast(dict[str, object], compared["diff"])
    semantic = cast(dict[str, object], diff["semantic"])
    normalized = cast(dict[str, object], semantic["normalized_ir"])
    assert "4100" in str(normalized["left"])
    assert "4800" in str(normalized["right"])
    assert isinstance(diff["diff_id"], str)
    assert compare_receipt.result_event_cursor is not None

    draft_receipt, drafted = _submit(
        service,
        "draft-report",
        "office.draft",
        {
            "template_artifact_id": imported["report-template.docx"]["artifact_id"],
            "diff_id": diff["diff_id"],
            "output_name": "comparison-report.docx",
        },
    )
    approval = cast(dict[str, object], drafted["approval"])
    approval_id = cast(str, approval["approval_id"])
    draft_id = cast(str, drafted["draft_id"])
    output = adapter.surface_authority.office.service.home / "artifacts" / "drafts" / "comparison-report.docx"
    assert not output.exists()
    pending = store.get_pending(approval_id)
    assert pending is not None
    pending_payload = cast(dict[str, object], pending["payload"])
    assert "template" not in pending_payload
    sealed_artifact = pending_payload.get("draft_artifact")
    sealed_uri = (
        cast(dict[str, object], sealed_artifact).get("uri")
        if isinstance(sealed_artifact, dict)
        else None
    )
    sealed_bytes = Path(sealed_uri).read_bytes() if isinstance(sealed_uri, str) else None

    events_before_answer = service.events()
    diff_event = next(event for event in events_before_answer if event.type == "office.diff_ready")
    requested = next(event for event in events_before_answer if event.type == "approval.requested")
    assert diff_event.command_id == "compare-spreadsheets"
    assert requested.command_id == "draft-report"
    assert requested.payload["approval_id"] == approval_id
    assert requested.payload["diff_id"] == diff["diff_id"]
    assert requested.payload["draft_id"] == draft_id
    assert diff_event.cursor < requested.cursor <= cast(int, draft_receipt.result_event_cursor)

    answer_receipt, answered = _submit(
        service,
        "approve-report-save",
        "approval.answer",
        {"approval_id": approval_id, "decision": "approve"},
    )
    saved = cast(dict[str, object], answered["artifact"])
    validation = cast(dict[str, object], answered["validation"])
    assert answered["outcome"] == "approved"
    assert output.is_file()
    assert Path(cast(str, saved["uri"])) == output
    assert validation["valid"] is True
    assert answer_receipt.result_event_cursor is not None

    with zipfile.ZipFile(output) as package:
        xml = "\n".join(
            package.read(name).decode("utf-8")
            for name in package.namelist()
            if name.endswith(".xml")
        )
    assert "BIRKIN_P3_03_DOCUMENT_SENTINEL" in xml
    assert "4100" in xml and "4800" in xml

    report = Document(str(output))
    assert [paragraph.text for paragraph in report.paragraphs[:3]] == [
        "Birkin Office Report",
        "BIRKIN_P3_03_DOCUMENT_SENTINEL",
        "기준값 4100, 후보값 4700",
    ]
    changes = [
        table
        for table in report.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["Field", "Old value", "New value"]
    ]
    assert len(changes) == 1
    rows = [[cell.text for cell in row.cells] for row in changes[0].rows[1:]]
    expected_entries = [
        (left, right)
        for left, right in zip(
            cast(list[dict[str, object]], normalized["left"]),
            cast(list[dict[str, object]], normalized["right"]),
            strict=True,
        )
        if left != right
    ]
    assert rows == [
        [
            f"{left['kind']} {left['order']}",
            cast(str, left["text"]),
            cast(str, right["text"]),
        ]
        for left, right in expected_entries
    ]
    assert all("4700" not in cell for row in rows for cell in row)
    assert any("4800" in cell for row in rows for cell in row)
    assert sealed_bytes is not None
    assert isinstance(sealed_uri, str)
    assert Path(sealed_uri).is_relative_to(
        adapter.surface_authority.office.service.home / "artifacts" / "staging"
    )
    assert output.read_bytes() == sealed_bytes

    _open_receipt, opened = _submit(
        service,
        "verify-saved-report",
        "office.open",
        {"artifact": saved},
    )
    assert cast(dict[str, object], opened["document"])["source"]

    answered_event = next(
        event for event in service.events() if event.type == "approval.answered"
    )
    activity_event = next(
        event for event in service.events() if event.type == "receipt.recorded"
    )
    assert answered_event.command_id == "approve-report-save"
    assert activity_event.command_id == "approve-report-save"
    assert activity_event.payload["approval_id"] == approval_id
    assert activity_event.payload["artifact_id"] == saved["artifact_id"]
    assert activity_event.payload["diff_id"] == diff["diff_id"]
    assert activity_event.payload["draft_id"] == draft_id
    assert activity_event.payload["request_command_id"] == "draft-report"
    assert activity_event.payload["approval_command_id"] == "approve-report-save"

    activity = next(
        panel.items
        for panel in service.snapshot().panels
        if panel.key == "activity_logs"
    )
    projected = next(item for item in activity if item["id"] == activity_event.event_id)
    assert projected["approval_id"] == approval_id
    assert projected["artifact_id"] == saved["artifact_id"]
    assert projected["request_command_id"] == "draft-report"


def test_rejection_never_saves_office_draft(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("BIRKIN_HOME", str(tmp_path / "home"))
    service, adapter = _service(tmp_path)
    artifacts: list[dict[str, object]] = []
    for index, name in enumerate(
        ("baseline.xlsx", "candidate.xlsx", "report-template.docx"), 1
    ):
        _receipt, result = _submit(
            service,
            f"reject-import-{index}",
            "file.import",
            {"source_path": str((FIXTURES / name).resolve())},
        )
        artifacts.append(_artifact(result))
    _receipt, compared = _submit(
        service,
        "reject-compare",
        "office.compare",
        {
            "left_artifact_id": artifacts[0]["artifact_id"],
            "right_artifact_id": artifacts[1]["artifact_id"],
        },
    )
    diff = cast(dict[str, object], compared["diff"])
    _receipt, drafted = _submit(
        service,
        "reject-draft",
        "office.draft",
        {
            "template_artifact_id": artifacts[2]["artifact_id"],
            "diff_id": diff["diff_id"],
            "output_name": "must-not-exist.docx",
        },
    )
    approval = cast(dict[str, object], drafted["approval"])

    _receipt, rejected = _submit(
        service,
        "reject-save",
        "approval.answer",
        {"approval_id": approval["approval_id"], "decision": "reject"},
    )

    assert rejected["outcome"] == "rejected"
    assert not (
        adapter.surface_authority.office.service.home
        / "artifacts"
        / "drafts"
        / "must-not-exist.docx"
    ).exists()
    assert all(event.type != "receipt.recorded" for event in service.events())
